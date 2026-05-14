# -*- coding: utf-8 -*-
"""Generuje raport końcowy projektu Kategoryzator w formacie .docx i .pdf."""

import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from _report_data import *  # noqa: F401,F403

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, Image
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── Rejestracja fontów z obsługą polskich znaków ──────────────────────────────
_FONT_PATHS = {
    "Calibri":        "C:/Windows/Fonts/calibri.ttf",
    "Calibri-Bold":   "C:/Windows/Fonts/calibrib.ttf",
    "Calibri-Italic": "C:/Windows/Fonts/calibrii.ttf",
}
_FONT_FALLBACK = {
    "Calibri":        "C:/Windows/Fonts/arial.ttf",
    "Calibri-Bold":   "C:/Windows/Fonts/arialbd.ttf",
    "Calibri-Italic": "C:/Windows/Fonts/ariali.ttf",
}
for _name, _path in _FONT_PATHS.items():
    _fp = _path if os.path.exists(_path) else _FONT_FALLBACK.get(_name, "")
    if _fp and os.path.exists(_fp):
        pdfmetrics.registerFont(TTFont(_name, _fp))

FONT_NORMAL = "Calibri"
FONT_BOLD   = "Calibri-Bold"
FONT_ITALIC = "Calibri-Italic"

# ── Kolory Symfonia ───────────────────────────────────────────────────────────
S_ORANGE      = "#E8521A"
S_DARK_ORANGE = "#B83E10"
S_LIGHT       = "#FEF0E8"
S_ALT_ROW     = "#FDDDC8"
S_DARK        = "#1A1A1A"
S_GRAY        = "#5A5A5A"

BASE_DIR    = os.path.join(HERE, "..")
PLOTS_DIR   = os.path.join(BASE_DIR, "plots")
OUTPUT_DIR  = BASE_DIR
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "raport_koncowy_kategoryzator.docx")
OUTPUT_PDF  = os.path.join(OUTPUT_DIR, "raport_koncowy_kategoryzator.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS – WORD
# ─────────────────────────────────────────────────────────────────────────────

def _shade_row(row, hex_color):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.replace("#", ""))
        tcPr.append(shd)


def _word_set_col_widths(table, widths_inches):
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_inches):
                cell.width = Inches(widths_inches[j])


def _add_word_table(doc, headers, rows,
                    header_bg="E8521A", alt_bg="FDDDC8",
                    col_widths=None, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style     = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = t.rows[0]
    _shade_row(hdr, header_bg)
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size      = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(rows):
        row = t.add_row()
        if r_idx % 2 == 1:
            _shade_row(row, alt_bg)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size = Pt(font_size)

    if col_widths:
        _word_set_col_widths(t, col_widths)
    return t


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS – PDF
# ─────────────────────────────────────────────────────────────────────────────

def _rl_h1():
    return ParagraphStyle("H1", fontName=FONT_BOLD, fontSize=16,
        textColor=colors.HexColor(S_ORANGE), spaceAfter=8, spaceBefore=18, leading=20)

def _rl_h2():
    return ParagraphStyle("H2", fontName=FONT_BOLD, fontSize=13,
        textColor=colors.HexColor(S_DARK_ORANGE), spaceAfter=5, spaceBefore=12, leading=16)

def _rl_h3():
    return ParagraphStyle("H3", fontName=FONT_BOLD, fontSize=11,
        textColor=colors.HexColor(S_DARK_ORANGE), spaceAfter=4, spaceBefore=10, leading=13)

def _rl_bd():
    return ParagraphStyle("Body", fontName=FONT_NORMAL, fontSize=9.5,
        textColor=colors.black, spaceAfter=6, leading=14, alignment=TA_JUSTIFY)

def _rl_blt():
    return ParagraphStyle("Bullet", fontName=FONT_NORMAL, fontSize=9.5,
        textColor=colors.black, spaceAfter=3, leading=13, leftIndent=14)

def _rl_sml():
    return ParagraphStyle("Small", fontName=FONT_NORMAL, fontSize=8,
        textColor=colors.HexColor(S_GRAY), spaceAfter=4, leading=11)

def _rl_ts():
    return TableStyle([
        ("FONTNAME",       (0, 0), (-1, 0),  FONT_BOLD),
        ("FONTSIZE",       (0, 0), (-1, 0),  9),
        ("BACKGROUND",     (0, 0), (-1, 0),  colors.HexColor(S_ORANGE)),
        ("TEXTCOLOR",      (0, 0), (-1, 0),  colors.white),
        ("ALIGN",          (0, 0), (-1, 0),  "CENTER"),
        ("FONTNAME",       (0, 1), (-1, -1), FONT_NORMAL),
        ("FONTSIZE",       (0, 1), (-1, -1), 8.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(S_ALT_ROW)]),
        ("GRID",           (0, 0), (-1, -1), 0.4, colors.grey),
        ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING",    (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 4),
        ("TOPPADDING",     (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 3),
    ])

def _hr(story, color=S_ORANGE, thick=1.5):
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=thick,
        color=colors.HexColor(color), spaceAfter=4))


def _img(path, w=14*cm):
    if os.path.exists(path):
        return Image(path, width=w, height=w * 0.6)
    return None


def _tbl(rows_data, widths):
    t = Table(rows_data, colWidths=widths)
    t.setStyle(_rl_ts())
    return t


# ─────────────────────────────────────────────────────────────────────────────
# WORD GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_word():
    doc = Document()

    for lvl, sz, rgb in [(1, 16, (0xE8, 0x52, 0x1A)),
                         (2, 14, (0xB8, 0x3E, 0x10)),
                         (3, 12, (0xB8, 0x3E, 0x10))]:
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name      = "Calibri"
        hs.font.size      = Pt(sz)
        hs.font.color.rgb = RGBColor(*rgb)
        hs.font.bold      = True

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Strona tytułowa ──────────────────────────────────────────────────────
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(TITLE)
    tr.bold = True; tr.font.size = Pt(36); tr.font.color.rgb = RGBColor(0xE8, 0x52, 0x1A)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(SUBTITLE)
    sr.font.size = Pt(14); sr.font.color.rgb = RGBColor(0xB8, 0x3E, 0x10)

    doc.add_paragraph()
    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(REPORT_HEADING)
    rr.bold = True; rr.font.size = Pt(18); rr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    doc.add_paragraph()
    mt = doc.add_table(rows=len(INFO_META), cols=2)
    mt.style = "Table Grid"
    for i, (k, v) in enumerate(INFO_META):
        mt.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        mt.rows[i].cells[1].paragraphs[0].add_run(v)
    doc.add_page_break()

    # ── Spis treści ──────────────────────────────────────────────────────────
    doc.add_heading("Spis Treści", level=1)
    for item in TOC_ITEMS:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ── 1. Streszczenie ──────────────────────────────────────────────────────
    doc.add_heading("1. Streszczenie wykonawcze", level=1)
    for txt in SEC_EXEC_SUMMARY:
        doc.add_paragraph(txt)
    doc.add_heading("Kluczowe wyniki", level=2)
    _add_word_table(doc, ["Obszar", "Wynik"], EXEC_BULLETS, col_widths=[2.0, 4.5])
    doc.add_page_break()

    # ── 2. Cel biznesowy ─────────────────────────────────────────────────────
    doc.add_heading("2. Cel i kontekst biznesowy projektu", level=1)
    doc.add_paragraph(SEC_BUSINESS)
    for a in ATTRS:
        doc.add_paragraph(a, style="List Bullet")
    doc.add_paragraph(SEC_BUSINESS_2)
    doc.add_page_break()

    # ── 3. Architektura ──────────────────────────────────────────────────────
    doc.add_heading("3. Architektura systemu", level=1)
    _add_word_table(doc, ["Warstwa", "Opis"], ARCH_ITEMS, col_widths=[2.2, 4.3])
    doc.add_heading("3.1 Flow decyzyjny predykcji", level=2)
    _add_word_table(doc, ["Krok", "Opis", "Status"],
                    [(s[0], s[1], s[2]) for s in FLOW_STEPS],
                    col_widths=[0.9, 4.0, 1.6])
    doc.add_page_break()

    # ── 4. Dane ──────────────────────────────────────────────────────────────
    doc.add_heading("4. Dane treningowe", level=1)
    doc.add_heading("4.1 Plik dane_ai.csv (profil KPIR)", level=2)
    _add_word_table(doc, ["Kolumna", "Opis"], KPIR_COLS, col_widths=[2.5, 4.0])
    doc.add_heading("4.2 Plik dane_ai_ryczalt.csv (profil ADVANCED)", level=2)
    doc.add_paragraph(SEC_DATA_2)
    doc.add_heading("4.3 Podzial Train/Test", level=2)
    doc.add_paragraph(SEC_DATA_SPLIT)
    doc.add_page_break()

    # ── 5. Feature Engineering ───────────────────────────────────────────────
    doc.add_heading("5. Metodologia i Feature Engineering", level=1)
    _add_word_table(doc, ["Cecha", "Opis"], FE_ITEMS, col_widths=[2.2, 4.3])
    doc.add_paragraph(SEC_FE_EXTRA)
    doc.add_page_break()

    # ── 6. Modele ML ─────────────────────────────────────────────────────────
    doc.add_heading("6. Modele uczenia maszynowego", level=1)
    doc.add_paragraph(SEC_ENSEMBLE)
    doc.add_heading("6.1 Keras – sieci neuronowe", level=2)
    _add_word_table(doc, ["Blok", "Opis"], KERAS_ARCH, col_widths=[2.2, 4.3])
    doc.add_heading("6.2 XGBoost – gradient boosting", level=2)
    _add_word_table(doc, ["Parametr", "Wartość"], XGB_PARAMS, col_widths=[2.5, 4.0])
    doc.add_heading("6.3 Hyperparameter Tuning (Keras Tuner)", level=2)
    doc.add_paragraph(
        "Tuning realizowany dla modelu cel_zakupu (profil ADVANCED) "
        "przy użyciu Keras Tuner – RandomSearch / BayesianOptimization.")
    _add_word_table(doc, ["Metryka/Parametr", "Wartość (best trial)"],
                    TUNING_DATA, col_widths=[2.5, 4.0])
    doc.add_page_break()

    # ── 7. Wyniki ─────────────────────────────────────────────────────────────
    doc.add_heading("7. Wyniki i metryki modeli", level=1)
    doc.add_heading("7.1 Wyniki XGBoost (profil KPIR)", level=2)
    _add_word_table(doc,
        ["Model (kolumna)", "Nazwa", "Accuracy", "Macro F1", "W-ROC AUC", "M-ROC AUC"],
        XGB_RESULTS, col_widths=[2.0, 2.0, 0.8, 0.8, 1.2, 1.2])
    doc.add_heading("7.2 Wyniki Keras (profil KPIR)", level=2)
    _add_word_table(doc,
        ["Model (kolumna)", "Nazwa", "Accuracy", "Obserwacje"],
        KERAS_RESULTS, col_widths=[2.0, 2.0, 0.8, 2.2])
    doc.add_heading("7.3 Profil ADVANCED (Ryczalt) – podsumowanie", level=2)
    _add_word_table(doc, ["Kolumna", "Obserwacje"], ADVANCED_SUMMARY,
                    col_widths=[2.5, 4.0])
    doc.add_page_break()

    # ── 8. Wykresy ───────────────────────────────────────────────────────────
    doc.add_heading("8. Wykresy z trenowania (Keras)", level=1)
    doc.add_paragraph(
        "Wykresy accuracy i loss dla każdej z sześciu klasyfikowanych kolumn "
        "(profil KPIR), wygenerowane podczas procesu treningu.")
    for idx, (key, (label, acc_f, loss_f)) in enumerate(PLOT_FILES.items()):
        doc.add_heading(f"8.{idx+1} {label}", level=2)
        for fname in (acc_f, loss_f):
            fpath = os.path.join(PLOTS_DIR, fname)
            if os.path.exists(fpath):
                doc.add_picture(fpath, width=Inches(5.5))
            else:
                doc.add_paragraph(f"[Wykres niedostępny: {fname}]")
    doc.add_page_break()

    # ── 9. API ───────────────────────────────────────────────────────────────
    doc.add_heading("9. Architektura API (FastAPI)", level=1)
    doc.add_heading("9.1 Endpointy REST", level=2)
    et = _add_word_table(doc, ["Metoda", "Ścieżka", "Opis"], ENDPOINTS,
                         col_widths=[0.9, 2.8, 3.0])
    for i, (method, _p, _d) in enumerate(ENDPOINTS):
        row = et.rows[i + 1]
        runs = row.cells[0].paragraphs[0].runs
        if runs:
            if method == "POST":
                runs[0].font.color.rgb = RGBColor(0xE8, 0x52, 0x1A)
            elif method == "GET":
                runs[0].font.color.rgb = RGBColor(0x6A, 0x30, 0x08)
            elif method in ("DELETE", "PUT"):
                runs[0].font.color.rgb = RGBColor(0x80, 0x30, 0x10)

    doc.add_heading("9.2 Pola żądania /predict", level=2)
    _add_word_table(doc, ["Pole", "Typ", "Opis"], PREDICT_FIELDS,
                    col_widths=[1.5, 1.3, 3.7])
    doc.add_heading("9.3 Fallback LLM (AWS Bedrock)", level=2)
    doc.add_paragraph(
        "Gdy similarity historii < confidence_ai_threshold i firma ma llm_enabled=True, "
        "system wywołuje AWS Bedrock Converse API. Prompt zawiera nazwę dokumentu, "
        "dostępne klasy i typ transakcji. Użycie tokenów jest rejestrowane w llm_usage.")
    doc.add_heading("9.4 Historia SQLite + FTS5", level=2)
    doc.add_paragraph(
        "Każde potwierdzone kategoryzowanie jest zapisywane w SQLite. "
        "FTS5 + cosine similarity; similarity >= confidence_exact → "
        "historical_match bez uruchamiania modeli ML.")
    doc.add_page_break()

    # ── 10. Infrastruktura ───────────────────────────────────────────────────
    doc.add_heading("10. Infrastruktura i wdrożenie", level=1)
    _add_word_table(doc, ["Komponent", "Szczegóły"], INFRA_ITEMS,
                    col_widths=[1.5, 5.0])
    doc.add_heading("10.1 Zmienne środowiskowe", level=2)
    env_items = [
        ("BOOKED_DB_PATH",           "Ścieżka do pliku SQLite historii"),
        ("DEFAULT_CONFIDENCE_EXACT", "Próg historii (domyślnie 0.90)"),
        ("DEFAULT_CONFIDENCE_AI",    "Próg modeli ML (domyślnie 0.70)"),
        ("LLM_BEDROCK_MODEL_ID",     "ID modelu AWS Bedrock"),
        ("SQS_QUEUE_URL",            "URL kolejki SQS"),
        ("AWS_ACCESS_KEY_ID",        "Klucz AWS (env / IAM role)"),
    ]
    _add_word_table(doc, ["Zmienna", "Opis"], env_items, col_widths=[2.5, 4.0])
    doc.add_page_break()

    # ── 11. Testy ─────────────────────────────────────────────────────────────
    doc.add_heading("11. Testy integracyjne", level=1)
    _add_word_table(doc, ["Plik testu", "Zakres"], TEST_FILES, col_widths=[2.8, 3.7])
    doc.add_paragraph("Uruchomienie: pytest tests/ -v")
    doc.add_page_break()

    # ── 12. Przypadki testowe ────────────────────────────────────────────────
    doc.add_heading("12. Przypadki testowe (UAT)", level=1)
    doc.add_paragraph("Szkielet do uzupełnienia podczas sesji UAT z księgowymi.")
    _add_word_table(doc, TC_HEADERS, [list(tc) for tc in MOCK_CASES],
                    col_widths=[0.6, 1.5, 1.5, 1.5, 1.2, 1.2, 0.7, 0.7], font_size=8)
    doc.add_page_break()

    # ── 13. Wnioski ──────────────────────────────────────────────────────────
    doc.add_heading("13. Wnioski końcowe", level=1)
    for title, text in CONCLUSIONS:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
    doc.add_page_break()

    # ── 14. Rekomendacje ─────────────────────────────────────────────────────
    doc.add_heading("14. Rekomendacje i dalszy rozwój", level=1)
    for rec in RECOMMENDATIONS:
        doc.add_paragraph(rec, style="List Bullet")
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fr = fp.add_run(FINAL_NOTE)
    fr.bold = True

    doc.save(OUTPUT_DOCX)
    print(f"[WORD] Zapisano: {OUTPUT_DOCX}")


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf():
    h1  = _rl_h1()
    h2  = _rl_h2()
    h3  = _rl_h3()
    bd  = _rl_bd()
    blt = _rl_blt()
    sml = _rl_sml()

    title_s = ParagraphStyle("TT", fontName=FONT_BOLD, fontSize=32,
        textColor=colors.HexColor(S_ORANGE), alignment=TA_CENTER, spaceAfter=10, leading=38)
    sub_s   = ParagraphStyle("TS", fontName=FONT_BOLD, fontSize=13,
        textColor=colors.HexColor(S_DARK_ORANGE), alignment=TA_CENTER, spaceAfter=6, leading=16)
    rh_s    = ParagraphStyle("TH", fontName=FONT_BOLD, fontSize=18,
        textColor=colors.HexColor(S_DARK), alignment=TA_CENTER, spaceAfter=20, leading=22)

    doc = SimpleDocTemplate(OUTPUT_PDF, pagesize=A4,
        rightMargin=2*cm, leftMargin=3*cm, topMargin=2.5*cm, bottomMargin=2.5*cm,
        title="Raport Końcowy – Kategoryzator")
    story = []

    # Strona tytułowa
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(TITLE, title_s))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(SUBTITLE, sub_s))
    story.append(Spacer(1, 0.5*cm))
    _hr(story)
    story.append(Paragraph(REPORT_HEADING, rh_s))
    story.append(Spacer(1, 0.5*cm))

    meta_rows = [[k, v] for k, v in INFO_META]
    mt = Table(meta_rows, colWidths=[4.5*cm, 11.5*cm])
    mt.setStyle(TableStyle([
        ("FONTNAME",       (0, 0), (0, -1), FONT_BOLD),
        ("FONTNAME",       (1, 0), (1, -1), FONT_NORMAL),
        ("FONTSIZE",       (0, 0), (-1, -1), 9),
        ("TEXTCOLOR",      (0, 0), (0, -1), colors.HexColor(S_ORANGE)),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor(S_LIGHT)]),
        ("LEFTPADDING",    (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 4),
        ("TOPPADDING",     (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 3),
    ]))
    story.append(mt)
    story.append(PageBreak())

    # Spis treści
    story.append(Paragraph("Spis Treści", h1))
    _hr(story, thick=0.8)
    for item in TOC_ITEMS:
        story.append(Paragraph(item, blt))
    story.append(PageBreak())

    # 1. Streszczenie
    story.append(Paragraph("1. Streszczenie wykonawcze", h1))
    _hr(story, thick=0.8)
    for txt in SEC_EXEC_SUMMARY:
        story.append(Paragraph(txt, bd))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("Kluczowe wyniki", h2))
    story.append(_tbl(
        [[Paragraph("<b>Obszar</b>", bd), Paragraph("<b>Wynik</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in EXEC_BULLETS],
        [5*cm, 11*cm]))
    story.append(PageBreak())

    # 2. Cel biznesowy
    story.append(Paragraph("2. Cel i kontekst biznesowy projektu", h1))
    _hr(story, thick=0.8)
    story.append(Paragraph(SEC_BUSINESS, bd))
    for a in ATTRS:
        story.append(Paragraph(f"• {a}", blt))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(SEC_BUSINESS_2, bd))
    story.append(PageBreak())

    # 3. Architektura
    story.append(Paragraph("3. Architektura systemu", h1))
    _hr(story, thick=0.8)
    story.append(_tbl(
        [[Paragraph("<b>Warstwa</b>", bd), Paragraph("<b>Opis</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in ARCH_ITEMS],
        [5*cm, 11*cm]))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph("3.1 Flow decyzyjny predykcji", h2))
    story.append(_tbl(
        [[Paragraph(f"<b>{h}</b>", bd) for h in ("Krok", "Opis", "Status")]] +
        [[Paragraph(s[0], bd), Paragraph(s[1], bd), Paragraph(s[2], bd)]
         for s in FLOW_STEPS],
        [2*cm, 9.5*cm, 4.5*cm]))
    story.append(PageBreak())

    # 4. Dane
    story.append(Paragraph("4. Dane treningowe", h1))
    _hr(story, thick=0.8)
    story.append(Paragraph("4.1 Plik dane_ai.csv (profil KPIR)", h2))
    story.append(_tbl(
        [[Paragraph("<b>Kolumna</b>", bd), Paragraph("<b>Opis</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in KPIR_COLS],
        [6*cm, 10*cm]))
    story.append(Paragraph("4.2 Plik dane_ai_ryczalt.csv (profil ADVANCED)", h2))
    story.append(Paragraph(SEC_DATA_2, bd))
    story.append(Paragraph("4.3 Podzial Train/Test", h2))
    story.append(Paragraph(SEC_DATA_SPLIT, bd))
    story.append(PageBreak())

    # 5. Feature Engineering
    story.append(Paragraph("5. Metodologia i Feature Engineering", h1))
    _hr(story, thick=0.8)
    story.append(_tbl(
        [[Paragraph("<b>Cecha</b>", bd), Paragraph("<b>Opis</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in FE_ITEMS],
        [5*cm, 11*cm]))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(SEC_FE_EXTRA, bd))
    story.append(PageBreak())

    # 6. Modele ML
    story.append(Paragraph("6. Modele uczenia maszynowego", h1))
    _hr(story, thick=0.8)
    story.append(Paragraph(SEC_ENSEMBLE, bd))
    story.append(Paragraph("6.1 Keras – sieci neuronowe", h2))
    story.append(_tbl(
        [[Paragraph("<b>Blok</b>", bd), Paragraph("<b>Opis</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in KERAS_ARCH],
        [5*cm, 11*cm]))
    story.append(Paragraph("6.2 XGBoost – gradient boosting", h2))
    story.append(_tbl(
        [[Paragraph("<b>Parametr</b>", bd), Paragraph("<b>Wartość</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in XGB_PARAMS],
        [6*cm, 10*cm]))
    story.append(Paragraph("6.3 Hyperparameter Tuning", h2))
    story.append(Paragraph(
        "Tuning realizowany dla modelu cel_zakupu (profil ADVANCED) "
        "przy użyciu Keras Tuner – RandomSearch / BayesianOptimization.", bd))
    story.append(_tbl(
        [[Paragraph("<b>Metryka/Parametr</b>", bd), Paragraph("<b>Wartość (best trial)</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in TUNING_DATA],
        [7*cm, 9*cm]))
    story.append(PageBreak())

    # 7. Wyniki
    story.append(Paragraph("7. Wyniki i metryki modeli", h1))
    _hr(story, thick=0.8)
    story.append(Paragraph("7.1 Wyniki XGBoost (profil KPIR)", h2))
    xr_hdr = ["Model", "Nazwa", "Accuracy", "Macro F1", "W-ROC AUC", "M-ROC AUC"]
    story.append(_tbl(
        [[Paragraph(f"<b>{h}</b>", bd) for h in xr_hdr]] +
        [[Paragraph(c, bd) for c in row] for row in XGB_RESULTS],
        [3.5*cm, 3.5*cm, 1.8*cm, 1.8*cm, 2.2*cm, 2.2*cm]))
    story.append(Paragraph("7.2 Wyniki Keras (profil KPIR)", h2))
    kr_hdr = ["Model", "Nazwa", "Accuracy", "Obserwacje"]
    story.append(_tbl(
        [[Paragraph(f"<b>{h}</b>", bd) for h in kr_hdr]] +
        [[Paragraph(c, bd) for c in row] for row in KERAS_RESULTS],
        [3.5*cm, 3.5*cm, 2*cm, 7*cm]))
    story.append(Paragraph("7.3 Profil ADVANCED (Ryczalt) – podsumowanie", h2))
    story.append(_tbl(
        [[Paragraph("<b>Kolumna</b>", bd), Paragraph("<b>Obserwacje</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in ADVANCED_SUMMARY],
        [6*cm, 10*cm]))
    story.append(PageBreak())

    # 8. Wykresy
    story.append(Paragraph("8. Wykresy z trenowania (Keras)", h1))
    _hr(story, thick=0.8)
    story.append(Paragraph(
        "Wykresy accuracy i loss dla każdej z sześciu kolumn (profil KPIR).", bd))
    story.append(Spacer(1, 0.3*cm))
    for idx, (key, (label, acc_f, loss_f)) in enumerate(PLOT_FILES.items()):
        story.append(Paragraph(f"8.{idx+1} {label}", h2))
        for fname in (acc_f, loss_f):
            fpath = os.path.join(PLOTS_DIR, fname)
            im = _img(fpath)
            if im:
                story.append(im)
                story.append(Paragraph(fname, sml))
            else:
                story.append(Paragraph(f"[Wykres niedostępny: {fname}]", sml))
        story.append(Spacer(1, 0.3*cm))
    story.append(PageBreak())

    # 9. API
    story.append(Paragraph("9. Architektura API (FastAPI)", h1))
    _hr(story, thick=0.8)
    story.append(Paragraph("9.1 Endpointy REST", h2))
    story.append(_tbl(
        [[Paragraph(f"<b>{h}</b>", bd) for h in ("Metoda", "Ścieżka", "Opis")]] +
        [[Paragraph(m, bd), Paragraph(p, bd), Paragraph(d, bd)]
         for m, p, d in ENDPOINTS],
        [2*cm, 6*cm, 8*cm]))
    story.append(Paragraph("9.2 Pola żądania /predict", h2))
    story.append(_tbl(
        [[Paragraph(f"<b>{h}</b>", bd) for h in ("Pole", "Typ", "Opis")]] +
        [[Paragraph(f, bd), Paragraph(t, bd), Paragraph(d, bd)]
         for f, t, d in PREDICT_FIELDS],
        [3.5*cm, 3.5*cm, 9*cm]))
    story.append(Paragraph("9.3 Fallback LLM (AWS Bedrock)", h2))
    story.append(Paragraph(
        "Gdy similarity historii < confidence_ai_threshold i firma ma llm_enabled=True, "
        "system wywołuje AWS Bedrock Converse API. Użycie tokenów rejestrowane w llm_usage.", bd))
    story.append(Paragraph("9.4 Historia SQLite + FTS5", h2))
    story.append(Paragraph(
        "FTS5 full-text search + cosine similarity. "
        "similarity >= confidence_exact → historical_match bez uruchamiania modeli ML.", bd))
    story.append(PageBreak())

    # 10. Infrastruktura
    story.append(Paragraph("10. Infrastruktura i wdrożenie", h1))
    _hr(story, thick=0.8)
    story.append(_tbl(
        [[Paragraph("<b>Komponent</b>", bd), Paragraph("<b>Szczegóły</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in INFRA_ITEMS],
        [4*cm, 12*cm]))
    story.append(Paragraph("10.1 Zmienne środowiskowe", h2))
    env2 = [
        ("BOOKED_DB_PATH",           "Ścieżka do pliku SQLite historii"),
        ("DEFAULT_CONFIDENCE_EXACT", "Próg historii (domyślnie 0.90)"),
        ("DEFAULT_CONFIDENCE_AI",    "Próg modeli ML (domyślnie 0.70)"),
        ("LLM_BEDROCK_MODEL_ID",     "ID modelu AWS Bedrock"),
        ("SQS_QUEUE_URL",            "URL kolejki SQS"),
        ("AWS_ACCESS_KEY_ID",        "Klucz AWS (env / IAM role)"),
    ]
    story.append(_tbl(
        [[Paragraph("<b>Zmienna</b>", bd), Paragraph("<b>Opis</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in env2],
        [6*cm, 10*cm]))
    story.append(PageBreak())

    # 11. Testy
    story.append(Paragraph("11. Testy integracyjne", h1))
    _hr(story, thick=0.8)
    story.append(_tbl(
        [[Paragraph("<b>Plik testu</b>", bd), Paragraph("<b>Zakres</b>", bd)]] +
        [[Paragraph(k, bd), Paragraph(v, bd)] for k, v in TEST_FILES],
        [6*cm, 10*cm]))
    story.append(Paragraph("Uruchomienie: <b>pytest tests/ -v</b>", bd))
    story.append(PageBreak())

    # 12. Przypadki testowe
    story.append(Paragraph("12. Przypadki testowe (UAT)", h1))
    _hr(story, thick=0.8)
    story.append(Paragraph("Szkielet do uzupełnienia podczas sesji UAT z księgowymi.", bd))
    story.append(_tbl(
        [[Paragraph(f"<b>{h}</b>", sml) for h in TC_HEADERS]] +
        [[Paragraph(str(c), sml) for c in tc] for tc in MOCK_CASES],
        [1.2*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2*cm, 2*cm, 1.2*cm, 2.1*cm]))
    story.append(PageBreak())

    # 13. Wnioski
    story.append(Paragraph("13. Wnioski końcowe", h1))
    _hr(story, thick=0.8)
    for title, text in CONCLUSIONS:
        story.append(Paragraph(title, h2))
        story.append(Paragraph(text, bd))
    story.append(PageBreak())

    # 14. Rekomendacje
    story.append(Paragraph("14. Rekomendacje i dalszy rozwój", h1))
    _hr(story, thick=0.8)
    for rec in RECOMMENDATIONS:
        story.append(Paragraph(f"• {rec}", blt))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(f"<b>{FINAL_NOTE}</b>", bd))

    doc.build(story)
    print(f"[PDF ] Zapisano: {OUTPUT_PDF}")


# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generowanie raportu końcowego Kategoryzator...")
    generate_word()
    generate_pdf()
    print("Gotowe!")
